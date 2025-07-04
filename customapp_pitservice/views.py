from django.http import JsonResponse
from rest_framework.decorators import api_view, permission_classes, renderer_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework.renderers import JSONRenderer
from django.shortcuts import render
from django.views.decorators.csrf import ensure_csrf_cookie
from django.views.decorators.csrf import csrf_exempt
import pdfkit
from commonapp.bixmodels.user_record import *
from commonapp.bixmodels.user_table import *
from commonapp.bixmodels.helper_db import *
from commonapp.helper import *
import locale
import re
from docx import Document
import os


@csrf_exempt
def stampa_bollettino(request):
    post_data = json.loads(request.body)
    data = {}
    

    recordid_bollettino = ''
    if request.method == 'POST':
        recordid_bollettino = post_data.get('recordid', '')

    record_bollettino = UserRecord('bollettini', recordid_bollettino)

    # Estrazione sicura degli ID
    recordid_stabile = record_bollettino.get_field('recordidstabile_').get('value', '') or ''
    recordid_dipendente = record_bollettino.get_field('recordiddipendente_').get('value', '') or ''
    recordid_cliente = record_bollettino.get_field('recordidcliente_').get('value', '') or ''

    # Creazione sicura dei record
    try:
        record_stabile = UserRecord('stabile', recordid_stabile) if recordid_stabile.strip() else None
        if record_stabile and record_stabile.values is None:
            record_stabile = None
    except:
        record_stabile = None

    try:
        record_dipendente = UserRecord('dipendente', recordid_dipendente) if recordid_dipendente.strip() else None
        if record_dipendente and record_dipendente.values is None: 
            record_dipendente = None
    except:
        record_dipendente = None

    try:
        record_cliente = UserRecord('cliente', recordid_cliente) if recordid_cliente.strip() else None
        if record_cliente and record_cliente.values is None:
            record_cliente = None
    except:
        record_cliente = None

    # Safe access ai campi
    def get_value_safe(record, fieldname):
        try:
            return record.get_field(fieldname).get('value', '') if record else ''
        except:
            return ''

    data['nome_cliente'] = get_value_safe(record_cliente, 'nome_cliente')
    data['indirizzo_cliente'] = get_value_safe(record_cliente, 'indirizzo')
    data['cap_cliente'] = get_value_safe(record_cliente, 'cap')
    data['citta_cliente'] = get_value_safe(record_cliente, 'citta')
    data['riferimento'] = get_value_safe(record_stabile, 'riferimento')
    data['citta_stabile'] = get_value_safe(record_stabile, 'citta')
    if get_value_safe(record_bollettino, 'data'):
        data['data'] = datetime.datetime.strptime(get_value_safe(record_bollettino, 'data'), "%Y-%m-%d").strftime("%d.%m.%Y")
    else:
        data['data'] = ''
    data['dipendente'] = f"{get_value_safe(record_dipendente, 'nome')} {get_value_safe(record_dipendente, 'cognome')}".strip()
    data['informazioni'] = get_value_safe(record_bollettino, 'informazioni')
    data['contattatoda'] = get_value_safe(record_bollettino, 'contattatoda')
    data['causa'] = get_value_safe(record_bollettino, 'causa')
    data['interventorichiesto'] = get_value_safe(record_bollettino, 'interventorichiesto')
    data['id'] = get_value_safe(record_bollettino, 'id')
    data['nr'] = get_value_safe(record_bollettino, 'nr')
    data['sostituzionedal'] = get_value_safe(record_bollettino, 'sostituzionedal')
    data['sostituzioneal'] = get_value_safe(record_bollettino, 'sostituzioneal')
    data['notelavoro'] = get_value_safe(record_bollettino, 'notelavoro')

    script_dir = os.path.dirname(os.path.abspath(__file__))
    wkhtmltopdf_path = script_dir + '\\wkhtmltopdf.exe'
    config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
    
    tipo_bollettino=record_bollettino.get_field('tipo_bollettino')['value']

    if tipo_bollettino=='Generico':
        content = render_to_string('pdf/bollettino_generico.html', data)
    if tipo_bollettino=='Sostituzione':
        content = render_to_string('pdf/bollettino_sostituzione.html', data)
    if tipo_bollettino=='Pulizia':
        content = render_to_string('pdf/bollettino_pulizia.html', data)
    if tipo_bollettino=='Tinteggio':
        content = render_to_string('pdf/bollettino_tinteggio.html', data)
    if tipo_bollettino=='Picchetto':
        content = render_to_string('pdf/bollettino_picchetto.html', data)
    if tipo_bollettino=='Giardino':
        content = render_to_string('pdf/bollettino_giardino.html', data)

    filename = f"Nr.{data['nr']} {record_stabile.values['indirizzo']} {tipo_bollettino} .pdf"
    if tipo_bollettino=='Sostituzione':
        filename = f"Nr.{data['nr']} {record_stabile.values['indirizzo']} Sostituzione Dal {data['sostituzionedal']} Al {data['sostituzioneal']}.pdf"
    filename = re.sub(r'[\\/*?:"<>|]', "", filename)
    filename_with_path = os.path.dirname(os.path.abspath(__file__))
    filename_with_path = filename_with_path.rsplit('views', 1)[0]
    filename_with_path = filename_with_path + '\\static\\pdf\\' + filename
    print(filename_with_path)
    #filename_with_path2 = os.path.join(settings.BASE_DIR, 'static', 'pdf', filename)
    #print(filename_with_path2)
    pdfkit.from_string(
    content,
    filename_with_path,
    configuration=config,
    options={
        "enable-local-file-access": "",
        # "quiet": ""  # <-- rimuovilo!
    }
    )


    #return HttpResponse(content)

    try:
        with open(filename_with_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/pdf")
            response['Content-Disposition'] = f'inline; filename="{filename}"'
            return response
        return response

    finally:
        os.remove(filename_with_path)


@csrf_exempt
def stampa_bollettino_test(request):
    data={}
    filename='test.pdf'
    

    content = render_to_string('pdf/bollettino_test.html', data)

    script_dir = os.path.dirname(os.path.abspath(__file__))
    wkhtmltopdf_path = script_dir + '\\wkhtmltopdf.exe'
    config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
    filename_with_path = os.path.dirname(os.path.abspath(__file__))
    filename_with_path = filename_with_path.rsplit('views', 1)[0]
    filename_with_path = filename_with_path + '\\static\\pdf\\' + filename
    pdfkit.from_string(
        content,
        filename_with_path,
        configuration=config,
        options={
            "enable-local-file-access": "",
            # "quiet": ""  # <-- rimuovilo!
        }
    )

    try:
        with open(filename_with_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/pdf")
            response['Content-Disposition'] = f'inline; filename={filename}'
            return response
        return response

    finally:
        os.remove(filename_with_path)



@csrf_exempt
def prepara_email(request):
    data = json.loads(request.body)
    tableid= data.get("tableid")
    recordid= data.get("recordid")
    type= data.get("type")
    print(tableid,recordid)

    email_fields = {
            "to": "",
            "cc": "",
            "bcc": "",	
            "subject": "",
            "text": "",
            "attachment_fullpath": "",
            "attachment_relativepath": "",
            "attachment_name": "",
            }
    
    if type == 'emailLavanderia':
        rendiconto_recordid=recordid
        rendiconto_record=UserRecord('rendicontolavanderia',rendiconto_recordid)
        mese=rendiconto_record.values['mese'][3:]
        anno=rendiconto_record.values['anno']
        stato=rendiconto_record.values['stato']
        stabile_recordid=rendiconto_record.values['recordidstabile_']
        stabile_record=UserRecord('stabile',stabile_recordid)
        stabile_riferimento=stabile_record.values['riferimento']
        stabile_indirizzo=stabile_record.values['indirizzo']
        stabile_citta=stabile_record.values['citta']
        sql=f"SELECT * FROM user_contattostabile WHERE deleted_='N' AND recordidstabile_='{stabile_recordid}'"
        row=HelpderDB.sql_query_row(sql)
        contatto_emai=''
        if row:
            contatto_recordid=row['recordidcontatti_']
            contatto_record=UserRecord('contatti',contatto_recordid)
            if contatto_record:
                contatto_emai=contatto_record.values['email']

        attachment_fullpath=HelpderDB.get_uploadedfile_fullpath('rendicontolavanderia',rendiconto_recordid,'allegato')
        attachment_relativepath=HelpderDB.get_uploadedfile_relativepath('rendicontolavanderia',rendiconto_recordid,'allegato')
        subject=f"Resoconto lavanderia - {stabile_riferimento} {stabile_citta} - {mese} {anno}"

        body = ""
        if stato=='Da fare':
            body = "Rendiconto da fare"

        if stato=='Inviato':
            body = "Rendiconto già inviato"

        if stato=='Preparato':
            body=f"""

                <p>
                    Egregi Signori,<br/>
                    Con la presente in allegato trasmettiamo il resoconto delle lavanderie dello stabile in {stabile_indirizzo} a {stabile_citta}.<br/>
                    Restiamo volentieri a disposizione e porgiamo cordiali saluti.
                </p>
                <br/>
                <table style="border: none; border-collapse: collapse; margin-top: 20px;">
                        <tr>
                            <td style="vertical-align: top; padding-right: 10px;">
                                <img src="https://pitservice.ch/wp-content/uploads/2025/04/miniminilogo.png" alt="Pit Service Logo">
                            </td>
                            <td style="font-family: Arial, sans-serif; font-size: 14px; line-height: 1.0">
                                <p>
                                    <b>Pit Service Sagl</b><br/>
                                    La cura del tuo immobile<br/>
                                    Phone: 091.993.03.92 <br/>
                                    Via San Gottardo 26 <br/>
                                    6943 Vezia <br/>
                                </p>
                            </td>
                        </tr>
                    </table>

                """
        
        if stato=='Nessuna ricarica':
            body=f"""
<p>
Egregi Signori,

con la presente per informarvi che durante il mese corrente non abbiamo eseguito ricariche tessere lavanderia presso lo stabile in {stabile_indirizzo} a {stabile_citta}.

Cordiali saluti
</p>
<br/>
<table style="border: none; border-collapse: collapse; margin-top: 20px;">
                    <tr>
                        <td style="vertical-align: top; padding-right: 10px;">
                            <img src="https://pitservice.ch/wp-content/uploads/2025/04/miniminilogo.png" alt="Pit Service Logo">
                        </td>
                        <td style="font-family: Arial, sans-serif; font-size: 14px; line-height: 1.0">
                            <p>
                                <b>Pit Service Sagl</b><br/>
                                La cura del tuo immobile<br/>
                                Phone: 091.993.03.92 <br/>
                                Via San Gottardo 26 <br/>
                                6943 Vezia <br/>
                            </p>
                        </td>
                    </tr>
                </table>
            """

        email_fields = {
            "to": contatto_emai,
            "cc": "contabilita@pitservice.ch,segreteria@pitservice.ch",
            "bcc": "",	
            "subject": subject,
            "text": body,
            "attachment_fullpath": attachment_fullpath,
            "attachment_relativepath": attachment_relativepath,
            "attachment_name": f"{stabile_riferimento} {stabile_citta} - Lavanderia - {mese} - {anno}.pdf",
            }
    
    if type == 'emailGasolio':
        stabile_recordid=recordid
        stabile_record=UserRecord('stabile',stabile_recordid)
        meseLettura='2025-04'
        anno, mese = meseLettura.split('-')

        sql=f"SELECT * FROM user_contattostabile WHERE deleted_='N' AND recordidstabile_='{stabile_recordid}'"
        row=HelpderDB.sql_query_row(sql)
        contatto_email=''
        if row:
            contatto_recordid=row['recordidcontatti_']
            contatto_record=UserRecord('contatti',contatto_recordid)
            if contatto_record:
                contatto_email=contatto_record.values['email']

        attachment_relativepath=stampa_gasoli(request)
        riferimento=stabile_record.values.get('riferimento', '')
        stabile_citta=stabile_record.values['citta']
        subject=f"Livello Gasolio - 05 {anno} - {riferimento} {stabile_citta}"
        body=f"""
         <p>
                Egregi Signori,<br/>
                con la presente in allegato trasmettiamo la lettura gasolio dello stabile in {stabile_record.values['indirizzo']}<br/>
                Restiamo volentieri a disposizione e porgiamo cordiali saluti.<br/>
        </p>

                <table style="border: none; border-collapse: collapse; margin-top: 20px;">
                    <tr>
                        <td style="vertical-align: top; padding-right: 10px;">
                            <img src="https://pitservice.ch/wp-content/uploads/2025/04/miniminilogo.png"  alt="Pit Service Logo">
                        </td>
                        <td style="font-family: Arial, sans-serif; font-size: 14px; ">
                            <p>
                                <b>Pit Service Sagl</b><br/>
                                La cura del tuo immobile<br/>
                                Phone: 091.993.03.92 <br/>
                                Via San Gottardo 26 <br/>
                                6943 Vezia <br/>
                            </p>
                        </td>
                    </tr>
                </table>
                """
        
        email_fields = {
            "to": contatto_email,
            "cc": "contabilita@pitservice.ch,segreteria@pitservice.ch",
            "bcc": "",	
            "subject": subject,
            "text": body,
            "attachment_fullpath": "",
            "attachment_relativepath": attachment_relativepath,
            "attachment_name": f"Lettura_Gasolio_05-{anno}-{riferimento}-{stabile_citta}.pdf",
            }

    return JsonResponse({"success": True, "emailFields": email_fields})

@csrf_exempt
def stampa_gasoli(request):
    data={}
    filename='report gasolio.pdf'
    recordid_stabile = ''
    data = json.loads(request.body)
    if request.method == 'POST':
        recordid_stabile = data.get('recordid')
        #meseLettura=data.get('date')
        #TODO sistemare dinamico
        meseLettura="2025 06-Giugno"
        anno, mese = meseLettura.split(' ')
    script_dir = os.path.dirname(os.path.abspath(__file__))
    wkhtmltopdf_path = script_dir + '\\wkhtmltopdf.exe'
    config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
    
    record_stabile=UserRecord('stabile',recordid_stabile)
    data['stabile']=record_stabile.values
    sql=f"""
    SELECT t.recordid_,t.anno,t.mese,t.datalettura,t.letturacm,t.letturalitri, i.riferimento, i.livellominimo, i.capienzacisterna, i.note
    FROM user_letturagasolio t
    INNER JOIN (
        SELECT recordidinformazionigasolio_, MAX(datalettura) AS max_datalettura
        FROM user_letturagasolio
        WHERE anno='{anno}' AND mese like '%{mese}%' AND deleted_='N' AND recordidstabile_ = '{recordid_stabile}'
        GROUP BY recordidinformazionigasolio_
        
    ) subquery
    ON t.recordidinformazionigasolio_ = subquery.recordidinformazionigasolio_ 
    AND t.datalettura = subquery.max_datalettura
    INNER JOIN user_informazionigasolio i
    ON t.recordidinformazionigasolio_ = i.recordid_
    WHERE t.recordidstabile_ = '{recordid_stabile}' AND t.deleted_ = 'N' 
            """
    ultimeletturegasolio = HelpderDB.sql_query(sql)
    data['ultimeletturegasolio']=ultimeletturegasolio
    data["show_letturacm"] = any(l.get('letturacm') for l in ultimeletturegasolio)
    data["show_note"] = any(l.get('note') for l in ultimeletturegasolio)

    content = render_to_string('pdf/gasolio.html', data)

    script_dir = os.path.dirname(os.path.abspath(__file__))
    wkhtmltopdf_path = script_dir + '\\wkhtmltopdf.exe'
    config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
    
    filename = f"Lettura Gasolio {mese} {anno}  {record_stabile.values['indirizzo']}.pdf"
    filename = re.sub(r'[\\/*?:"<>|]', "", filename)
    #filename='gasolio.pdf'

    filename_with_path = os.path.dirname(os.path.abspath(__file__))
    filename_with_path = filename_with_path.rsplit('views', 1)[0]
    filename_with_path = filename_with_path + '\\static\\pdf\\' + filename
    pdfkit.from_string(
        content,
        filename_with_path,
        configuration=config,
        options={
            "enable-local-file-access": "",
            # "quiet": ""  # <-- rimuovilo!
        }
    )

    if False:
        return 'customapp_pitservice/static/pdf/' + filename
    else:
        try:
            with open(filename_with_path, 'rb') as fh:
                response = HttpResponse(fh.read(), content_type="application/pdf")
                response['Content-Disposition'] = f'inline; filename={filename}'
                return response
            return response

        finally:
            os.remove(filename_with_path)



def crea_lista_lavanderie(request):
    # Extract any needed data from the request, if relevant
    # mese = request.POST.get('mese')
    # Ottieni la data attuale
    # Imposta la localizzazione in italiano
    post_data = json.loads(request.body)
    meserichiesto=  post_data.get('mese')
    locale.setlocale(locale.LC_TIME, "it_IT.utf8")

    now = datetime.datetime.now()
    current_year=now.year
    current_month_num = now.month
    
    # Mese corrente
    if meserichiesto == 'mesecorrente':
        month = current_month_num # Formato numerico (es. 2)
        year=current_year
        month_2digit = f"{month:02d}"  # Due cifre (es. "02")
        month_name = now.strftime("%B").capitalize()  # Nome esteso con prima lettera maiuscola (es. "Febbraio")
    else:
        # Mese successivo
        month = current_month_num + 1 if current_month_num < 12 else 1  # Gestisce il passaggio da dicembre a gennaio
        year = current_year  if current_month_num < 12 else current_year + 1
        month_2digit = f"{month:02d}"  # Due cifre (es. "03")
        month_name = datetime.date(now.year if month > 1 else now.year + 1, month, 1).strftime("%B").capitalize()

    mese=month_2digit+'-'+month_name
    sql=f"""
            SELECT DISTINCT recordidstabile_
            FROM user_informazionilavanderia
            WHERE deleted_ = 'N'
            AND recordidstabile_ NOT IN (
                SELECT recordidstabile_
                FROM user_rendicontolavanderia
                WHERE anno = '{year}' 
                AND mese = '{mese}'
                AND deleted_ = 'N'
            );


    """
    informazionilavanderia_records=HelpderDB.sql_query(sql)
    counter=0
    for informazionelavanderia in informazionilavanderia_records:
        record_rendiconto=UserRecord('rendicontolavanderia')
        record_rendiconto.values['recordidstabile_']=informazionelavanderia['recordidstabile_']
        record_stabile=UserRecord('stabile',informazionelavanderia['recordidstabile_'])
        record_rendiconto.values['recordidcliente_']=record_stabile.values['recordidcliente_']
        record_rendiconto.values['mese']=mese
        record_rendiconto.values['anno']=year
        record_rendiconto.values['stato']="Da fare"
        record_rendiconto.save()
        counter=counter+1
    # Return them in a JSON response (or use as needed)
    return JsonResponse({
        'success': True,
        'counter': counter
    })

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
def rimuovi_sezione(doc, inizio_marker, fine_marker):
    body = doc.element.body
    remove = False
    elements_to_remove = []

    for element in list(body):
        text = ''
        # Prova a leggere il testo del paragrafo
        if element.tag == qn('w:p'):
            text = ''.join(node.text or '' for node in element.iter() if node.tag == qn('w:t'))
        # Prova a leggere il testo della tabella (nel caso il marker sia in una cella)
        elif element.tag == qn('w:tbl'):
            for row in element.iter():
                if row.tag == qn('w:t'):
                    text += row.text or ''

        if inizio_marker in text:
            remove = True
            elements_to_remove.append(element)
            continue

        if remove:
            elements_to_remove.append(element)

        if fine_marker in text:
            remove = False
            continue

    # Rimuovi tutti gli elementi trovati
    for el in elements_to_remove:
        body.remove(el)


def replace_text_in_paragraph(paragraph, key, value):
    """Sostituisce il testo mantenendo lo stile originale"""
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


def download_offerta(request):
    data = json.loads(request.body)
    recordid = data.get('recordid')

    record = UserRecord('offerta', recordid)
    templateofferta = record.values.get('tipoofferta', '')

    filename = f"Offerta_{record.values.get('id', '')}.docx"
    
    script_dir = os.path.dirname(os.path.abspath(__file__))


    if templateofferta == 'Custodia':
        file_path = os.path.join(script_dir, 'static', 'template_offerte', 'servizio_custodia.docx')
        doc = Document(file_path)

        id_offerta = record.values.get('id', '')

        recordid_cliente = record.values.get('recordidcliente_', '')
        record_cliente = UserRecord('cliente', recordid_cliente)
        nome_cliente = record_cliente.values.get('nome_cliente', '')
        indirizzo_cliente = record_cliente.values.get('indirizzo', '')
        cap_cliente = record_cliente.values.get('cap', '')
        citta_cliente = record_cliente.values.get('citta', '')

        recordid_stabile = record.values.get('recordidstabile_', '')
        record_stabile = UserRecord('stabile', recordid_stabile)
        nome_stabile = record_stabile.values.get('titolo_stabile', '')
        indirizzo_stabile = record_stabile.values.get('indirizzo', '')
        citta_stabile = record_stabile.values.get('citta', '')

        data = datetime.datetime.now().strftime("%d.%m.%Y")

        # Ottieni le sezioni da mostrare dal database
        sections_to_show = record.values.get('sezionitemplate', '')
        sections_to_show = sections_to_show.split(';') if sections_to_show else []

        # Tutte le possibli sezioni nel template
        all_sections = ['custodia_pulizia', 'piscina', 'area_verde', 'neve']

        # Rimuovi le sezioni che NON sono presenti in sections_to_show
        for section in all_sections:
            if section not in sections_to_show:
                if section == 'custodia_pulizia':
                    rimuovi_sezione(doc, '{{inizio_custodia}}', '{{fine_custodia}}')
                    rimuovi_sezione(doc, '{{inizio_prezzo_custodia}}', '{{fine_prezzo_custodia}}')
                elif section == 'piscina':
                    rimuovi_sezione(doc, '{{inizio_piscina}}', '{{fine_piscina}}')
                    rimuovi_sezione(doc, '{{inizio_prezzo_piscina}}', '{{fine_prezzo_piscina}}')
                elif section == 'area_verde':
                    rimuovi_sezione(doc, '{{inizio_area_verde}}', '{{fine_area_verde}}')
                    rimuovi_sezione(doc, '{{inizio_prezzo_area_verde}}', '{{fine_prezzo_area_verde}}')
                elif section == 'neve':
                    rimuovi_sezione(doc, '{{inizio_prezzo_neve}}', '{{fine_prezzo_neve}}')
                    rimuovi_sezione(doc, '{{inizio_prezzo_neve}}', '{{fine_prezzo_neve}}')


        replacements = {
            '{{nome_cliente}}': nome_cliente,
            '{{indirizzo_cliente}}': indirizzo_cliente,
            '{{cap_cliente}}': cap_cliente,
            '{{citta_cliente}}': citta_cliente,
            '{{data}}': data,
            '{{id_offerta}}': str(id_offerta),
            '{{nome_stabile}}': nome_stabile,
            '{{indirizzo_stabile}}': indirizzo_stabile,
            '{{citta_stabile}}': citta_stabile,
            '{{inizio_custodia}}': '',
            '{{fine_custodia}}': '',
            '{{inizio_piscina}}': '',
            '{{fine_piscina}}': '',
            '{{inizio_area_verde}}': '',
            '{{fine_area_verde}}': '',
            '{{inizio_prezzo_custodia}}': '',
            '{{fine_prezzo_custodia}}': '',
            '{{inizio_prezzo_piscina}}': '',
            '{{fine_prezzo_piscina}}': '',
            '{{inizio_prezzo_area_verde}}': '',
            '{{fine_prezzo_area_verde}}': '',
            '{{inizio_prezzo_neve}}': '',
            '{{fine_prezzo_neve}}': '',
        }

        for p in doc.paragraphs:
            for key, value in replacements.items():
                replace_text_in_paragraph(p, key, value)

    elif templateofferta == 'Manutenzione giardino':
        file_path = os.path.join(script_dir, 'static', 'template_offerte', 'manutenzione_giardino.docx')
        doc = Document(file_path)

        id_offerta = record.values.get('id', '')

        recordid_cliente = record.values.get('recordidcliente_', '')
        record_cliente = UserRecord('cliente', recordid_cliente)
        nome_cliente = record_cliente.values.get('nome_cliente', '')
        indirizzo_cliente = record_cliente.values.get('indirizzo', '')
        cap_cliente = record_cliente.values.get('cap', '')
        citta_cliente = record_cliente.values.get('citta', '')

        recordid_stabile = record.values.get('recordidstabile_', '')
        record_stabile = UserRecord('stabile', recordid_stabile)
        indirizzo_stabile = record_stabile.values.get('indirizzo', '')
        citta_stabile = record_stabile.values.get('citta', '')

        data = datetime.datetime.now().strftime("%d.%m.%Y")

        replacements = {
            '{{nome_cliente}}': nome_cliente,
            '{{indirizzo_cliente}}': indirizzo_cliente,
            '{{cap_cliente}}': cap_cliente,
            '{{citta_cliente}}': citta_cliente,
            '{{id_offerta}}': str(id_offerta),
            '{{indirizzo_stabile}}': indirizzo_stabile,
            '{{citta_stabile}}': citta_stabile,
            '{{data}}': data
        }

        for p in doc.paragraphs:
            for key, value in replacements.items():
                replace_text_in_paragraph(p, key, value)
    
    elif templateofferta == 'Pulizia appartamento':
        file_path = os.path.join(script_dir, 'static', 'template_offerte', 'pulizia_appartamento.docx')
        doc = Document(file_path)

        id_offerta = record.values.get('id', '')
        importo = record.values.get('importo', '')

        recordid_cliente = record.values.get('recordidcliente_', '')
        record_cliente = UserRecord('cliente', recordid_cliente)
        nome_cliente = record_cliente.values.get('nome_cliente', '')
        indirizzo_cliente = record_cliente.values.get('indirizzo', '')
        cap_cliente = record_cliente.values.get('cap', '')
        citta_cliente = record_cliente.values.get('citta', '')

        recordid_stabile = record.values.get('recordidstabile_', '')
        record_stabile = UserRecord('stabile', recordid_stabile)
        indirizzo_stabile = record_stabile.values.get('indirizzo', '')
        citta_stabile = record_stabile.values.get('citta', '')

        data = datetime.datetime.now().strftime("%d.%m.%Y")

        replacements = {
            '{{nome_cliente}}': nome_cliente,
            '{{indirizzo_cliente}}': indirizzo_cliente,
            '{{cap_cliente}}': cap_cliente,
            '{{citta_cliente}}': citta_cliente,
            '{{id_offerta}}': str(id_offerta),
            '{{indirizzo_stabile}}': indirizzo_stabile,
            '{{citta_stabile}}': citta_stabile,
            '{{importo}}': importo,
            '{{data}}': data
        }

        for p in doc.paragraphs:
            for key, value in replacements.items():
                replace_text_in_paragraph(p, key, value)


    elif templateofferta == 'Tinteggio appartamento':
        file_path = os.path.join(script_dir, 'static', 'template_offerte', 'tinteggio_appartamento.docx')
        doc = Document(file_path)

        id_offerta = record.values.get('id', '')
        importo = record.values.get('importo', '')

        recordid_cliente = record.values.get('recordidcliente_', '')
        record_cliente = UserRecord('cliente', recordid_cliente)
        nome_cliente = record_cliente.values.get('nome_cliente', '')
        indirizzo_cliente = record_cliente.values.get('indirizzo', '')
        cap_cliente = record_cliente.values.get('cap', '')
        citta_cliente = record_cliente.values.get('citta', '')

        recordid_stabile = record.values.get('recordidstabile_', '')
        record_stabile = UserRecord('stabile', recordid_stabile)
        indirizzo_stabile = record_stabile.values.get('indirizzo', '')
        citta_stabile = record_stabile.values.get('citta', '')

        data = datetime.datetime.now().strftime("%d.%m.%Y")

        replacements = {
            '{{nome_cliente}}': nome_cliente,
            '{{indirizzo_cliente}}': indirizzo_cliente,
            '{{cap_cliente}}': cap_cliente,
            '{{citta_cliente}}': citta_cliente,
            '{{id_offerta}}': str(id_offerta),
            '{{indirizzo_stabile}}': indirizzo_stabile,
            '{{citta_stabile}}': citta_stabile,
            '{{importo}}': importo,
            '{{data}}': data
        }

        for p in doc.paragraphs:
            for key, value in replacements.items():
                replace_text_in_paragraph(p, key, value)

    modified_file_path = os.path.join(script_dir, 'static', 'modified_template.docx')
    doc.save(modified_file_path)

    if os.path.exists(modified_file_path):
        with open(modified_file_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/pdf")
            response['Content-Disposition'] = f'inline; filename={filename}'
        os.remove(modified_file_path)
        return response
    else:
        return JsonResponse({'error': 'File not found'}, status=404)
